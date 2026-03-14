"""
Telegram bot — long-polling receiver.

Only messages from the configured TELEGRAM_CHAT_ID are processed.
All other senders receive no response (silent drop).
"""

from __future__ import annotations

import asyncio
import contextlib
import html as _html
import logging
import os
import secrets
import shutil
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from urllib.parse import urlparse

from aiogram import BaseMiddleware, Bot, Dispatcher, F
from aiogram.filters import Command
from aiogram.types import (
    BotCommand,
    CallbackQuery,
    FSInputFile,
    InlineKeyboardButton,
    InlineKeyboardMarkup,
    Message,
    ReplyKeyboardRemove,
    Update,
)

from .config import Settings
from . import engine
from . import watchers_config as wc

# ── clipboard helpers ────────────────────────────────────────────────────────

_CLIPBOARD_BACKENDS = [
    ("windows",   ["/mnt/c/Windows/System32/clip.exe"]),
    ("xclip",     ["xclip", "-selection", "clipboard"]),
    ("xsel",      ["xsel", "--clipboard", "--input"]),
]


async def _set_clipboard(text: str) -> list[str]:
    """Write *text* to every available clipboard backend. Returns names that succeeded."""
    ok: list[str] = []
    for name, cmd in _CLIPBOARD_BACKENDS:
        try:
            proc = await asyncio.create_subprocess_exec(
                *cmd,
                stdin=asyncio.subprocess.PIPE,
                stdout=asyncio.subprocess.DEVNULL,
                stderr=asyncio.subprocess.DEVNULL,
            )
            await proc.communicate(text.encode())
            if proc.returncode == 0:
                ok.append(name)
        except (FileNotFoundError, OSError):
            pass
        except Exception:
            log.exception("clipboard backend %r failed", name)
    return ok

log = logging.getLogger("watcher.bot")


# ── pending action tracker ──────────────────────────────────────────────────────
# Maps chat_id → current pending context dict.
# "action" key is one of: "edit_interval" | "edit_prompt_item" | "create_prompt"
# Extra keys: watcher_id, prompt_idx, ask_msg_id, return_to

_pending: dict[int, dict] = {}

# ── prompt UI state ─────────────────────────────────────────────────────────────
# Maps chat_id → {wid, prompt_msg_ids: list[int], add_msg_id: int}
# Tracks the per-prompt messages currently shown so they can be cleaned up.

_prompts_ui: dict[int, dict] = {}


# ── files storage ───────────────────────────────────────────────────────────────

_FILES_DIR = Path.home() / ".config" / "watcher" / "files"


def _saved_files() -> list[Path]:
    if not _FILES_DIR.exists():
        return []
    return sorted(f for f in _FILES_DIR.iterdir() if f.is_file())


def _files_list_text(files: list[Path]) -> str:
    if not files:
        return "📁 <b>Files</b>\n\nNo files saved yet.\nSend any photo or document here to save it."
    n = len(files)
    return f"📁 <b>Files</b>  <i>({n} file{'s' if n != 1 else ''})</i>"


def _files_list_kb(files: list[Path]) -> InlineKeyboardMarkup:
    rows = [
        [
            InlineKeyboardButton(text=f"📄  {f.name}", callback_data=f"f:resend:{i}"),
            InlineKeyboardButton(text="🗑", callback_data=f"f:del:{i}"),
        ]
        for i, f in enumerate(files)
    ]
    rows.append([_done_btn()])
    return InlineKeyboardMarkup(inline_keyboard=rows)


# ── interval unit helpers ──────────────────────────────────────────────────────

_UNIT_LABELS: dict[str, str] = {"s": "Seconds", "m": "Minutes", "h": "Hours", "d": "Days"}
_UNIT_MULT:   dict[str, int] = {"s": 1, "m": 60, "h": 3600, "d": 86400}


# ── keyboard / text builders ───────────────────────────────────────────────────

def _done_btn() -> InlineKeyboardButton:
    return InlineKeyboardButton(text="✖ Done", callback_data="w:done")


def _fmt_interval(seconds: int) -> str:
    if seconds % 86400 == 0:
        return f"{seconds // 86400}d"
    if seconds % 3600 == 0:
        return f"{seconds // 3600}h"
    if seconds % 60 == 0:
        return f"{seconds // 60}m"
    return f"{seconds}s"


def _watchers_list_text(watchers: list[wc.WatcherConfig]) -> str:
    if not watchers:
        return "📡 <b>Watchers</b>\n\nNo watchers configured."
    active = sum(1 for w in watchers if w.enabled)
    paused = len(watchers) - active
    if paused == 0:
        subtitle = f"{active} active"
    elif active == 0:
        subtitle = f"{paused} paused"
    else:
        subtitle = f"{active} active · {paused} paused"
    return f"📡 <b>Watchers</b>  <i>({subtitle})</i>"


def _name_from_url(url: str) -> str:
    """Deduce a short display name from a URL."""
    try:
        p = urlparse(url)
        host = p.netloc
        if host.startswith("www."):
            host = host[4:]
        path = p.path.rstrip("/")
        return (host + path) if path else host
    except Exception:
        return url


def _watchers_list_kb(watchers: list[wc.WatcherConfig]) -> InlineKeyboardMarkup:
    rows = [
        [InlineKeyboardButton(
            text=f"{'🟢' if w.enabled else '🔴'}  {w.name}  ·  {_fmt_interval(w.interval)}",
            callback_data=f"w:actions:{w.id}",
        )]
        for w in watchers
    ]
    rows.append([InlineKeyboardButton(text="➕ Add", callback_data="w:add_watcher")])
    rows.append([_done_btn()])
    return InlineKeyboardMarkup(inline_keyboard=rows)


def _cancel_new_watcher_kb() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[[
        InlineKeyboardButton(text="✖ Cancel", callback_data="w:cancel_new_watcher"),
    ]])


def _skip_prompt_kb() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[[
        InlineKeyboardButton(text="⏭ Skip", callback_data="w:skip_prompt"),
        InlineKeyboardButton(text="✖ Cancel", callback_data="w:cancel_new_watcher"),
    ]])


def _unit_sel_existing_kb(wid: str) -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="Seconds", callback_data=f"w:iunit:{wid}:s"),
         InlineKeyboardButton(text="Minutes", callback_data=f"w:iunit:{wid}:m")],
        [InlineKeyboardButton(text="Hours",   callback_data=f"w:iunit:{wid}:h"),
         InlineKeyboardButton(text="Days",    callback_data=f"w:iunit:{wid}:d")],
        [InlineKeyboardButton(text="✖ Cancel", callback_data=f"w:cancel_input:{wid}")],
    ])


def _unit_sel_new_kb() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="Seconds", callback_data="w:nwiunit:s"),
         InlineKeyboardButton(text="Minutes", callback_data="w:nwiunit:m")],
        [InlineKeyboardButton(text="Hours",   callback_data="w:nwiunit:h"),
         InlineKeyboardButton(text="Days",    callback_data="w:nwiunit:d")],
        [InlineKeyboardButton(text="✖ Cancel", callback_data="w:cancel_new_watcher")],
    ])


def _actions_kb(w: wc.WatcherConfig) -> InlineKeyboardMarkup:
    toggle = "🔴 Disable" if w.enabled else "🟢 Enable"
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="⚡ Fetch Now",      callback_data=f"w:fetch_now:{w.id}")],
        [InlineKeyboardButton(text="📋 Get Snapshot",   callback_data=f"w:get_snapshot:{w.id}")],
        [InlineKeyboardButton(text=toggle,              callback_data=f"w:toggle:{w.id}")],
        [InlineKeyboardButton(text="✏️ Rename",         callback_data=f"w:rename:{w.id}")],
        [InlineKeyboardButton(text="⏱  Set Interval",  callback_data=f"w:interval:{w.id}")],
        [InlineKeyboardButton(text="📝 Prompts",        callback_data=f"w:prompts:{w.id}")],
        [InlineKeyboardButton(text="🗑 Delete",         callback_data=f"w:delete:{w.id}")],
        [InlineKeyboardButton(text="◀ Back",  callback_data="w:list"),
         _done_btn()],
    ])


def _prompt_kb(wid: str, idx: int) -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[[
        InlineKeyboardButton(text="🗑 Delete", callback_data=f"w:del_prompt:{wid}:{idx}"),
        InlineKeyboardButton(text="✏️ Modify", callback_data=f"w:edit_prompt:{wid}:{idx}"),
    ]])


def _add_prompt_kb(wid: str) -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="➕ Add", callback_data=f"w:add_prompt:{wid}"),
         InlineKeyboardButton(text="◀ Back",        callback_data=f"w:actions:{wid}")],
        [_done_btn()],
    ])


def _input_cancel_kb(wid: str) -> InlineKeyboardMarkup:
    """Keyboard attached to bot 'ask for input' messages so the user can cancel."""
    return InlineKeyboardMarkup(inline_keyboard=[[
        InlineKeyboardButton(text="✖  Done", callback_data=f"w:cancel_input:{wid}"),
    ]])


async def _cleanup_prompts_ui(bot: Bot, chat_id: int) -> None:
    """Delete all tracked prompt UI messages for chat_id."""
    ui = _prompts_ui.pop(chat_id, None)
    if not ui:
        return
    for mid in ui.get("prompt_msg_ids", []):
        with contextlib.suppress(Exception):
            await bot.delete_message(chat_id, mid)
    if ui.get("add_msg_id"):
        with contextlib.suppress(Exception):
            await bot.delete_message(chat_id, ui["add_msg_id"])


async def _render_prompts(bot: Bot, chat_id: int, wid: str) -> None:
    """Send one message per prompt then an ➕ Add / ◀ Back message. Cleans up old UI first."""
    await _cleanup_prompts_ui(bot, chat_id)

    w = wc.get(wid)
    if w is None:
        return

    total = len(w.prompts)
    prompt_msg_ids: list[int] = []
    for i, p in enumerate(w.prompts):
        header = f"<b>Prompt {i + 1} / {total}</b>\n\n"
        msg = await bot.send_message(
            chat_id, header + _html.escape(p),
            reply_markup=_prompt_kb(wid, i),
            parse_mode="HTML",
        )
        prompt_msg_ids.append(msg.message_id)

    name = _html.escape(w.name)
    summary = f"📝 <b>{name}</b> — {len(w.prompts)} prompt(s)" if w.prompts else f"📝 <b>{name}</b> — no prompts yet."
    add_msg = await bot.send_message(chat_id, summary, reply_markup=_add_prompt_kb(wid), parse_mode="HTML")

    _prompts_ui[chat_id] = {
        "wid": wid,
        "prompt_msg_ids": prompt_msg_ids,
        "add_msg_id": add_msg.message_id,
    }


def _watcher_info_text(w: wc.WatcherConfig) -> str:
    status = "🟢 enabled" if w.enabled else "🔴 disabled"
    return (
        f"<b>{_html.escape(w.name)}</b>\n"
        f"{_html.escape(w.url)}\n"
        f"Interval: {w.interval}s | {status}"
    )


# ── dispatcher ─────────────────────────────────────────────────────────────────

class _ChatGuard(BaseMiddleware):
    """Outer middleware — silently drops every update not from the configured chat."""

    def __init__(self, allowed_chat_id: int) -> None:
        self._allowed = allowed_chat_id

    async def __call__(self, handler, event: Update, data: dict):  # type: ignore[override]
        chat_id: int | None = None
        if event.message:
            chat_id = event.message.chat.id
        elif event.callback_query and event.callback_query.message:
            chat_id = event.callback_query.message.chat.id
        elif event.edited_message:
            chat_id = event.edited_message.chat.id
        elif event.channel_post:
            chat_id = event.channel_post.chat.id

        if chat_id is None or chat_id != self._allowed:
            if chat_id is not None:
                log.warning("blocked update from unauthorized chat_id=%d", chat_id)
            return

        return await handler(event, data)


async def _edit_to_actions(query: CallbackQuery, w: wc.WatcherConfig, note: str = "") -> None:
    """Edit the current message to the watcher actions view, with an optional status note."""
    text = _watcher_info_text(w) + (f"\n{note}" if note else "")
    await query.message.edit_text(text, reply_markup=_actions_kb(w), parse_mode="HTML")  # type: ignore[union-attr]


async def _convert_via_libreoffice(pdf_path: Path, tmp: str) -> Path:
    """Run LibreOffice headless conversion and return the .docx path."""
    log.debug("pdf2docx libreoffice: starting conversion input=%s outdir=%s", pdf_path, tmp)
    # Each conversion gets its own LO profile dir to avoid lock contention
    profile_dir = Path(tmp) / "lo_profile"
    profile_dir.mkdir()
    lo_env = {**os.environ, "LANG": "en_US.UTF-8", "LC_ALL": "en_US.UTF-8"}
    proc = await asyncio.create_subprocess_exec(
        "libreoffice",
        f"-env:UserInstallation=file://{profile_dir}",
        "--headless",
        "--norestore",
        "--infilter=writer_pdf_import",
        "--convert-to", "docx:MS Word 2007 XML",
        "--outdir", tmp, str(pdf_path),
        stdout=asyncio.subprocess.PIPE,
        stderr=asyncio.subprocess.PIPE,
        env=lo_env,
    )
    stdout, stderr = await proc.communicate()
    log.debug(
        "pdf2docx libreoffice: returncode=%d stdout=%s stderr=%s",
        proc.returncode, stdout.decode().strip(), stderr.decode().strip(),
    )
    if proc.returncode != 0:
        raise RuntimeError(stdout.decode() or stderr.decode() or f"exit code {proc.returncode}")
    matches = [p for p in Path(tmp).glob("*.docx") if p.is_file()]
    log.debug("pdf2docx libreoffice: output files found=%s", [m.name for m in matches])
    if not matches:
        raise RuntimeError("LibreOffice produced no output file.")
    log.debug("pdf2docx libreoffice: done output=%s", matches[0])
    return matches[0]


def _convert_via_pdf2docx(pdf_path: Path) -> Path:
    """Convert using the pdf2docx library (fallback when LibreOffice is absent)."""
    try:
        from pdf2docx import Converter  # type: ignore[import]
    except ImportError:
        raise RuntimeError(
            "Neither LibreOffice nor pdf2docx is available. "
            "Install LibreOffice for best results."
        )
    docx_path = pdf_path.with_suffix(".docx")
    log.debug("pdf2docx pdf2docx: starting conversion input=%s output=%s", pdf_path, docx_path)
    cv = Converter(str(pdf_path))
    cv.convert(str(docx_path))
    cv.close()
    log.debug("pdf2docx pdf2docx: conversion finished exists=%s", docx_path.exists())
    if not docx_path.exists():
        raise RuntimeError("pdf2docx produced no output file.")
    log.debug("pdf2docx pdf2docx: done output=%s", docx_path)
    return docx_path


def _docx_text_seems_garbled(docx_path: Path) -> bool:
    """Return True if the DOCX has no text or is dominated by replacement/private-use chars."""
    try:
        from docx import Document  # type: ignore[import]
    except ImportError:
        return False
    text = "".join(p.text for p in Document(str(docx_path)).paragraphs)
    if not text.strip():
        log.debug("pdf2docx garble-check: no text found in output")
        return True
    # U+FFFD = replacement char; U+E000-U+F8FF = Private Use Area (raw glyph IDs leak here)
    garbage = sum(1 for c in text if c == "\ufffd" or "\ue000" <= c <= "\uf8ff")
    ratio = garbage / len(text)
    log.debug("pdf2docx garble-check: garbage ratio=%.2f", ratio)
    return ratio > 0.05


def _convert_via_pymupdf(pdf_path: Path) -> Path:
    """Extract text with pymupdf and rebuild as DOCX — best for Hebrew/RTL and custom encodings."""
    try:
        import fitz  # type: ignore[import]
    except ImportError:
        raise RuntimeError("pymupdf is not installed. Run: pip install pymupdf")
    try:
        from docx import Document  # type: ignore[import]
        from docx.oxml import OxmlElement  # type: ignore[import]
    except ImportError:
        raise RuntimeError("python-docx is not installed. Run: pip install python-docx")

    def _is_rtl(text: str) -> bool:
        return any("\u0590" <= c <= "\u05ff" or "\u0600" <= c <= "\u06ff" or "\ufb1d" <= c <= "\ufb4f" for c in text)

    pdf_doc = fitz.open(str(pdf_path))  # type: ignore[call-overload]
    word_doc = Document()
    log.debug("pdf2docx pymupdf: starting conversion pages=%d", pdf_doc.page_count)
    for page_num, page in enumerate(pdf_doc):
        for block in page.get_text("blocks", sort=True):  # type: ignore[call-arg]
            if block[6] != 0:  # skip image blocks
                continue
            text = block[4].strip()
            if not text:
                continue
            para = word_doc.add_paragraph(text)
            if _is_rtl(text):
                pPr = para._p.get_or_add_pPr()
                pPr.append(OxmlElement("w:bidi"))
        if page_num < pdf_doc.page_count - 1:
            word_doc.add_page_break()
    pdf_doc.close()
    docx_path = pdf_path.with_suffix(".docx")
    word_doc.save(str(docx_path))
    log.debug("pdf2docx pymupdf: done output=%s", docx_path)
    return docx_path


async def _convert_via_online(pdf_path: Path, tmp: str, *, headed: bool = False) -> Path:
    """Convert PDF→DOCX using ilovepdf.com via Playwright (best for Hebrew/RTL)."""
    import re as _re
    from playwright.async_api import async_playwright, TimeoutError as PWTimeout

    tool_url = "https://www.ilovepdf.com/pdf_to_word"
    out_dir = Path(tmp)
    out_dir.mkdir(parents=True, exist_ok=True)

    log.debug("pdf2docx: online: navigating to ilovepdf")
    async with async_playwright() as pw:
        browser = await pw.chromium.launch(headless=not headed)
        context = await browser.new_context(accept_downloads=True)
        page = await context.new_page()

        try:
            # ── 1. Open the tool page ──────────────────────────────────────
            await page.goto(tool_url, wait_until="networkidle")

            # ── 2. Upload via the hidden <input type="file"> ───────────────
            file_input = page.locator('input[type="file"]').first
            await file_input.set_input_files(str(pdf_path))
            log.debug("pdf2docx: online: file set, waiting for upload to finish")

            await page.wait_for_selector(
                f'text="{pdf_path.name}"', timeout=30_000
            )

            # ── 3. Click the Convert button ────────────────────────────────
            convert_btn = page.get_by_role(
                "button", name=_re.compile(r"convert", _re.IGNORECASE)
            )
            await convert_btn.wait_for(state="visible", timeout=15_000)
            await convert_btn.click()
            log.debug("pdf2docx: online: conversion started")

            # ── 4. Wait for the Download link and trigger download ─────────
            download_link = page.get_by_role(
                "link", name=_re.compile(r"download", _re.IGNORECASE)
            )
            try:
                await download_link.wait_for(state="visible", timeout=120_000)
            except PWTimeout:
                raise RuntimeError(
                    "Download link never appeared – conversion may have failed."
                )

            async with page.expect_download(timeout=60_000) as dl_info:
                await download_link.click()

            download = await dl_info.value
            failure = await download.failure()
            if failure:
                raise RuntimeError(f"download failed: {failure}")
            suggested = download.suggested_filename or (pdf_path.stem + ".docx")
            out = out_dir / suggested
            await download.save_as(str(out))
            log.debug("pdf2docx: online: done output=%s", out)
            return out

        finally:
            await browser.close()


def _build_dispatcher(chat_id: int, settings: Settings) -> Dispatcher:
    dp = Dispatcher()
    dp.update.outer_middleware(_ChatGuard(chat_id))

    # ── Commands ───────────────────────────────────────────────────────────────

    @dp.message(Command("start"))
    async def cmd_start(message: Message) -> None:
        log.info("cmd=start chat_id=%d user=%s", message.chat.id, message.from_user.username if message.from_user else None)
        _pending.pop(message.chat.id, None)
        await message.answer(
            "Watcher is running.\n\nCommands:\n/status — service status\n/watchers — manage watchers\n/clipboard — set clipboard\n/help — help",
            reply_markup=ReplyKeyboardRemove(),
        )

    @dp.message(Command("status"))
    async def cmd_status(message: Message) -> None:
        log.info("cmd=status chat_id=%d user=%s", message.chat.id, message.from_user.username if message.from_user else None)
        await message.answer("Status: running")

    @dp.message(Command("help"))
    async def cmd_help(message: Message) -> None:
        log.info("cmd=help chat_id=%d user=%s", message.chat.id, message.from_user.username if message.from_user else None)
        await message.answer(
            "/start     — greeting\n"
            "/status    — service status\n"
            "/watchers  — manage watchers\n"
            "/clipboard — set clipboard\n"
            "/pdf2docx  — convert PDF → DOCX\n"
            "/help      — this message"
        )

    @dp.message(Command("clipboard"))
    async def cmd_clipboard(message: Message) -> None:
        log.info("cmd=clipboard chat_id=%d user=%s", message.chat.id, message.from_user.username if message.from_user else None)
        _pending.pop(message.chat.id, None)
        # Text may be passed inline: /clipboard some text here
        text = (message.text or "").partition(" ")[2].strip()
        if text:
            ok = await _set_clipboard(text)
            if ok:
                await message.answer(f"✅ Clipboard set via: {', '.join(ok)}")
            else:
                await message.answer("❌ No clipboard backend available (clip.exe / xclip / xsel).")
        else:
            ask = await message.answer("Send the text to copy to the clipboard:")
            _pending[message.chat.id] = {"action": "clipboard", "ask_msg_id": ask.message_id}

    @dp.message(Command("watchers"))
    async def cmd_watchers(message: Message) -> None:
        log.info("cmd=watchers chat_id=%d user=%s", message.chat.id, message.from_user.username if message.from_user else None)
        _pending.pop(message.chat.id, None)
        watchers = wc.load_all()
        await message.answer(
            _watchers_list_text(watchers),
            reply_markup=_watchers_list_kb(watchers),
            parse_mode="HTML",
        )

    # ── PDF → DOCX conversion ──────────────────────────────────────────────────

    @dp.message(Command("pdf2docx"))
    async def cmd_pdf2docx(message: Message) -> None:
        log.info("cmd=pdf2docx chat_id=%d", message.chat.id)
        _pending.pop(message.chat.id, None)
        ask = await message.answer("Send me a PDF file to convert to DOCX:")
        _pending[message.chat.id] = {"action": "pdf2docx", "ask_msg_id": ask.message_id}

    # ── Files ──────────────────────────────────────────────────────────────────

    @dp.message(Command("files"))
    async def cmd_files(message: Message) -> None:
        log.info("cmd=files chat_id=%d", message.chat.id)
        _FILES_DIR.mkdir(parents=True, exist_ok=True)
        files = _saved_files()
        await message.answer(
            _files_list_text(files),
            reply_markup=_files_list_kb(files),
            parse_mode="HTML",
        )

    @dp.callback_query(F.data == "f:list")
    async def cb_files_list(query: CallbackQuery) -> None:
        files = _saved_files()
        await query.message.edit_text(  # type: ignore[union-attr]
            _files_list_text(files),
            reply_markup=_files_list_kb(files),
            parse_mode="HTML",
        )
        await query.answer()

    @dp.callback_query(F.data.startswith("f:del:"))
    async def cb_file_delete(query: CallbackQuery) -> None:
        idx = int(query.data.split(":", 2)[2])  # type: ignore[union-attr]
        files = _saved_files()
        if idx >= len(files):
            await query.answer("File not found.", show_alert=True)
            return
        name = files[idx].name
        files[idx].unlink()
        await query.answer(f"🗑 {name} deleted.")
        remaining = _saved_files()
        await query.message.edit_text(  # type: ignore[union-attr]
            _files_list_text(remaining),
            reply_markup=_files_list_kb(remaining),
            parse_mode="HTML",
        )

    @dp.callback_query(F.data.startswith("f:resend:"))
    async def cb_file_resend(query: CallbackQuery, bot: Bot) -> None:
        idx = int(query.data.split(":", 2)[2])  # type: ignore[union-attr]
        files = _saved_files()
        if idx >= len(files):
            await query.answer("File not found.", show_alert=True)
            return
        await bot.send_document(query.message.chat.id, FSInputFile(files[idx]))  # type: ignore[union-attr]
        await query.answer()

    @dp.message(F.photo | F.document)
    async def handle_file(message: Message, bot: Bot) -> None:
        pending = _pending.get(message.chat.id)

        # ── pdf2docx conversion ────────────────────────────────────────────────
        if pending and pending.get("action") == "pdf2docx" and message.document:
            doc = message.document
            if doc.mime_type != "application/pdf":
                await message.answer("❌ That's not a PDF. Please send a .pdf file:")
                return
            _pending.pop(message.chat.id, None)
            progress = await message.answer("Converting… ⏳")
            with tempfile.TemporaryDirectory() as tmp:
                # Use a plain ASCII filename — LibreOffice fails on non-ASCII paths
                pdf_path = Path(tmp) / "input.pdf"
                original_stem = Path(doc.file_name or "output").stem
                await bot.download(doc.file_id, destination=pdf_path)
                try:
                    docx_path = await _convert_via_online(pdf_path, tmp, headed=settings.headed)
                    # try:
                        # docx_path = await _convert_via_online(pdf_path, tmp, headed=settings.headed)
                    # except Exception as exc:
                        # log.warning("pdf2doc: online conversion failed (%s), falling back to LibreOffice", exc)
                        # if shutil.which("libreoffice"):
                            # try:
                                # docx_path = await _convert_via_libreoffice(pdf_path, tmp)
                                # if _docx_text_seems_garbled(docx_path):
                                    # log.debug("pdf2doc: libreoffice output garbled, trying pymupdf")
                                    # docx_path = _convert_via_pymupdf(pdf_path)
                            # except Exception:
                                # docx_path = _convert_via_pdf2docx(pdf_path)
                        # else:
                            # docx_path = _convert_via_pdf2docx(pdf_path)
                except RuntimeError as exc:
                    await progress.edit_text(
                        f"❌ Conversion failed:\n<code>{_html.escape(str(exc))}</code>",
                        parse_mode="HTML",
                    )
                    return
                # Rename output to preserve the original filename for the recipient
                named_docx = docx_path.with_name(original_stem + ".docx")
                docx_path.rename(named_docx)
                await progress.edit_text("✅ Done")
                await bot.send_document(message.chat.id, FSInputFile(named_docx))
            return

        # ── existing file-save logic ──────────────────────────────────────────
        _FILES_DIR.mkdir(parents=True, exist_ok=True)

        if message.photo:
            file_id = message.photo[-1].file_id
            ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
            filename = f"photo_{ts}.jpg"
        else:
            file_id = message.document.file_id  # type: ignore[union-attr]
            filename = f"file_{datetime.now(timezone.utc).strftime('%Y%m%d_%H%M%S')}.doc"  # type: ignore[union-attr]

        dest = _FILES_DIR / filename
        if dest.exists():
            stem, suffix, counter = dest.stem, dest.suffix, 1
            while dest.exists():
                dest = _FILES_DIR / f"{stem}_{counter}{suffix}"
                counter += 1

        await bot.download(file_id, destination=dest)
        log.info("file saved: %s", dest)
        await message.answer(
            f"✅ Saved as <code>{_html.escape(dest.name)}</code>",
            parse_mode="HTML",
        )

    # ── Done — close session at any stage ──────────────────────────────────────

    @dp.callback_query(F.data == "w:done")
    async def cb_done(query: CallbackQuery, bot: Bot) -> None:
        chat_id = query.message.chat.id  # type: ignore[union-attr]
        _pending.pop(chat_id, None)
        await _cleanup_prompts_ui(bot, chat_id)
        with contextlib.suppress(Exception):
            await query.message.delete()  # type: ignore[union-attr]
        await query.answer("Done.")

    # ── Cancel pending input ────────────────────────────────────────────────────

    @dp.callback_query(F.data.startswith("w:cancel_input:"))
    async def cb_cancel_input(query: CallbackQuery, bot: Bot) -> None:
        chat_id = query.message.chat.id  # type: ignore[union-attr]
        wid = query.data.split(":", 2)[2]  # type: ignore[union-attr]
        pending = _pending.pop(chat_id, None)
        # Delete the ask message (the one carrying this Cancel button).
        with contextlib.suppress(Exception):
            await query.message.delete()  # type: ignore[union-attr]
        await query.answer("Cancelled.")
        return_to = (pending or {}).get("return_to", "prompts")
        if return_to == "prompts":
            await _render_prompts(bot, chat_id, wid)
        else:
            # return_to == "actions"
            w = wc.get(wid)
            if w:
                await bot.send_message(
                    chat_id,
                    _watcher_info_text(w),
                    reply_markup=_actions_kb(w),
                    parse_mode="HTML",
                )

    # ── Watcher list ───────────────────────────────────────────────────────────

    @dp.callback_query(F.data == "w:list")
    async def cb_list(query: CallbackQuery, bot: Bot) -> None:
        chat_id = query.message.chat.id  # type: ignore[union-attr]
        _pending.pop(chat_id, None)
        await _cleanup_prompts_ui(bot, chat_id)
        watchers = wc.load_all()
        await query.message.edit_text(  # type: ignore[union-attr]
            _watchers_list_text(watchers),
            reply_markup=_watchers_list_kb(watchers),
            parse_mode="HTML",
        )
        await query.answer()

    # ── Watcher actions ────────────────────────────────────────────────────────

    @dp.callback_query(F.data.startswith("w:actions:"))
    async def cb_actions(query: CallbackQuery, bot: Bot) -> None:
        chat_id = query.message.chat.id  # type: ignore[union-attr]
        _pending.pop(chat_id, None)
        wid = query.data.split(":", 2)[2]  # type: ignore[union-attr]
        w = wc.get(wid)
        if w is None:
            await query.answer("Watcher not found.", show_alert=True)
            return
        # Clean up individual prompt messages if navigating back from prompts view.
        ui = _prompts_ui.pop(chat_id, None)
        if ui:
            for mid in ui.get("prompt_msg_ids", []):
                with contextlib.suppress(Exception):
                    await bot.delete_message(chat_id, mid)
        await _edit_to_actions(query, w)
        await query.answer()

    @dp.callback_query(F.data.startswith("w:toggle:"))
    async def cb_toggle(query: CallbackQuery) -> None:
        wid = query.data.split(":", 2)[2]  # type: ignore[union-attr]
        w = wc.get(wid)
        if w is None:
            await query.answer("Watcher not found.", show_alert=True)
            return
        w.enabled = not w.enabled
        wc.save(w)
        await _edit_to_actions(query, w)
        await query.answer("Enabled." if w.enabled else "Disabled.")

    @dp.callback_query(F.data.startswith("w:fetch_now:"))
    async def cb_fetch_now(query: CallbackQuery) -> None:
        wid = query.data.split(":", 2)[2]  # type: ignore[union-attr]
        w = wc.get(wid)
        if w is None:
            await query.answer("Watcher not found.", show_alert=True)
            return
        await query.answer("⚡ Fetching…")
        await query.message.edit_text(  # type: ignore[union-attr]
            f"{_watcher_info_text(w)}\n⏳ Fetching…",
            parse_mode="HTML",
        )
        status = await engine.fetch_once(settings, w)
        w = wc.get(wid) or w
        notes = {
            "changed": "✅ Change detected — notification sent.",
            "ok":      "✅ No change.",
            "error":   "❌ Fetch failed — see logs.",
        }
        await _edit_to_actions(query, w, notes.get(status, status))

    @dp.callback_query(F.data.startswith("w:get_snapshot:"))
    async def cb_get_snapshot(query: CallbackQuery, bot: Bot) -> None:
        wid = query.data.split(":", 2)[2]  # type: ignore[union-attr]
        w = wc.get(wid)
        if w is None:
            await query.answer("Watcher not found.", show_alert=True)
            return
        _, text = engine._get_snapshot(wid)
        if text is None:
            await query.answer("No snapshot yet.", show_alert=True)
            return
        await query.answer()
        chat_id = query.message.chat.id  # type: ignore[union-attr]
        header = f"📋 <b>Snapshot: {_html.escape(w.name)}</b>\n\n"
        # Telegram message limit is 4096 chars; send as file if too large.
        MAX_MSG = 4000
        truncated = len(text) > MAX_MSG
        body = _html.escape(text[:MAX_MSG])
        suffix = f"\n…<i>(truncated, {len(text)} chars total)</i>" if truncated else ""
        await bot.send_message(chat_id, header + body + suffix, parse_mode="HTML")

    @dp.callback_query(F.data.startswith("w:delete:"))
    async def cb_delete(query: CallbackQuery) -> None:
        _pending.pop(query.message.chat.id, None)  # type: ignore[union-attr]
        wid = query.data.split(":", 2)[2]  # type: ignore[union-attr]
        w = wc.get(wid)
        name = _html.escape(w.name) if w else wid
        wc.delete(wid)
        remaining = wc.load_all()
        if remaining:
            await query.message.edit_text(  # type: ignore[union-attr]
                f"🗑 <b>{name}</b> deleted.\n\n{_watchers_list_text(remaining)}",
                reply_markup=_watchers_list_kb(remaining),
                parse_mode="HTML",
            )
        else:
            await query.message.edit_text(  # type: ignore[union-attr]
                f"🗑 <b>{name}</b> deleted.\n\nNo watchers remaining.",
                parse_mode="HTML",
            )
        await query.answer("Deleted.")

    # ── Interval editing ───────────────────────────────────────────────────────

    @dp.callback_query(F.data.startswith("w:interval:"))
    async def cb_interval(query: CallbackQuery) -> None:
        wid = query.data.split(":", 2)[2]  # type: ignore[union-attr]
        w = wc.get(wid)
        if w is None:
            await query.answer("Watcher not found.", show_alert=True)
            return
        await query.message.delete()  # type: ignore[union-attr]
        ask = await query.message.answer(  # type: ignore[union-attr]
            f"Current interval: <b>{w.interval}s</b>\n\nChoose the unit for the new interval:",
            reply_markup=_unit_sel_existing_kb(wid),
            parse_mode="HTML",
        )
        _pending[query.message.chat.id] = {  # type: ignore[union-attr]
            "action": "edit_interval_unit",
            "watcher_id": wid,
            "ask_msg_id": ask.message_id,
            "return_to": "actions",
        }
        await query.answer()

    @dp.callback_query(F.data.startswith("w:iunit:"))
    async def cb_iunit(query: CallbackQuery) -> None:
        parts = query.data.split(":")  # type: ignore[union-attr]
        wid, unit = parts[2], parts[3]
        if wc.get(wid) is None:
            await query.answer("Watcher not found.", show_alert=True)
            return
        label = _UNIT_LABELS[unit]
        await query.message.delete()  # type: ignore[union-attr]
        ask = await query.message.answer(  # type: ignore[union-attr]
            f"Send the new interval in {label.lower()}:",
            reply_markup=_input_cancel_kb(wid),
        )
        _pending[query.message.chat.id] = {  # type: ignore[union-attr]
            "action": "edit_interval",
            "watcher_id": wid,
            "unit": unit,
            "ask_msg_id": ask.message_id,
            "return_to": "actions",
        }
        await query.answer()

    # ── Rename ─────────────────────────────────────────────────────────────────

    @dp.callback_query(F.data.startswith("w:rename:"))
    async def cb_rename(query: CallbackQuery) -> None:
        wid = query.data.split(":", 2)[2]  # type: ignore[union-attr]
        w = wc.get(wid)
        if w is None:
            await query.answer("Watcher not found.", show_alert=True)
            return
        await query.message.delete()  # type: ignore[union-attr]
        ask = await query.message.answer(  # type: ignore[union-attr]
            f"Current name:\n\n<code>{_html.escape(w.name)}</code>\n\nSend the new name:",
            reply_markup=_input_cancel_kb(wid),
            parse_mode="HTML",
        )
        _pending[query.message.chat.id] = {  # type: ignore[union-attr]
            "action": "edit_name",
            "watcher_id": wid,
            "ask_msg_id": ask.message_id,
            "return_to": "actions",
        }
        await query.answer()

    # ── Prompts management ─────────────────────────────────────────────────────

    @dp.callback_query(F.data.startswith("w:prompts:"))
    async def cb_prompts(query: CallbackQuery, bot: Bot) -> None:
        wid = query.data.split(":", 2)[2]  # type: ignore[union-attr]
        if wc.get(wid) is None:
            await query.answer("Watcher not found.", show_alert=True)
            return
        _pending.pop(query.message.chat.id, None)  # type: ignore[union-attr]
        await query.message.delete()  # type: ignore[union-attr]
        await _render_prompts(bot, query.message.chat.id, wid)  # type: ignore[union-attr]
        await query.answer()

    @dp.callback_query(F.data.startswith("w:del_prompt:"))
    async def cb_del_prompt(query: CallbackQuery, bot: Bot) -> None:
        parts = query.data.split(":")  # type: ignore[union-attr]
        wid, idx = parts[2], int(parts[3])
        w = wc.get(wid)
        if w is None:
            await query.answer("Watcher not found.", show_alert=True)
            return
        if idx < 0 or idx >= len(w.prompts):
            await query.answer("Prompt already removed.", show_alert=True)
            return
        w.prompts.pop(idx)
        wc.save(w)
        await query.message.delete()  # type: ignore[union-attr]
        await query.answer("🗑 Deleted.")
        await _render_prompts(bot, query.message.chat.id, wid)  # type: ignore[union-attr]

    @dp.callback_query(F.data.startswith("w:edit_prompt:"))
    async def cb_edit_prompt(query: CallbackQuery) -> None:
        parts = query.data.split(":")  # type: ignore[union-attr]
        wid, idx = parts[2], int(parts[3])
        w = wc.get(wid)
        if w is None:
            await query.answer("Watcher not found.", show_alert=True)
            return
        if idx < 0 or idx >= len(w.prompts):
            await query.answer("Prompt not found.", show_alert=True)
            return
        await query.message.delete()  # type: ignore[union-attr]
        ask = await query.message.answer(  # type: ignore[union-attr]
            f"Current prompt {idx + 1}:\n\n<code>{_html.escape(w.prompts[idx])}</code>\n\nSend the new text:",
            reply_markup=_input_cancel_kb(wid),
            parse_mode="HTML",
        )
        _pending[query.message.chat.id] = {  # type: ignore[union-attr]
            "action": "edit_prompt_item",
            "watcher_id": wid,
            "prompt_idx": idx,
            "ask_msg_id": ask.message_id,
            "return_to": "prompts",
        }
        await query.answer()

    @dp.callback_query(F.data.startswith("w:add_prompt:"))
    async def cb_add_prompt(query: CallbackQuery) -> None:
        wid = query.data.split(":", 2)[2]  # type: ignore[union-attr]
        if wc.get(wid) is None:
            await query.answer("Watcher not found.", show_alert=True)
            return
        await query.message.delete()  # type: ignore[union-attr]
        ask = await query.message.answer(  # type: ignore[union-attr]
            "Send the new prompt text:",
            reply_markup=_input_cancel_kb(wid),
        )
        _pending[query.message.chat.id] = {  # type: ignore[union-attr]
            "action": "create_prompt",
            "watcher_id": wid,
            "ask_msg_id": ask.message_id,
            "return_to": "prompts",
        }
        await query.answer()

    # ── Add watcher flow ───────────────────────────────────────────────────────

    @dp.callback_query(F.data == "w:add_watcher")
    async def cb_add_watcher(query: CallbackQuery) -> None:
        chat_id = query.message.chat.id  # type: ignore[union-attr]
        _pending.pop(chat_id, None)
        await query.message.delete()  # type: ignore[union-attr]
        ask = await query.message.answer(  # type: ignore[union-attr]
            "Send the URL to watch:",
            reply_markup=_cancel_new_watcher_kb(),
        )
        _pending[chat_id] = {"action": "new_watcher_url", "ask_msg_id": ask.message_id}
        await query.answer()

    @dp.callback_query(F.data == "w:cancel_new_watcher")
    async def cb_cancel_new_watcher(query: CallbackQuery, bot: Bot) -> None:
        chat_id = query.message.chat.id  # type: ignore[union-attr]
        _pending.pop(chat_id, None)
        with contextlib.suppress(Exception):
            await query.message.delete()  # type: ignore[union-attr]
        await query.answer("Cancelled.")
        watchers = wc.load_all()
        await bot.send_message(
            chat_id,
            _watchers_list_text(watchers),
            reply_markup=_watchers_list_kb(watchers),
            parse_mode="HTML",
        )

    @dp.callback_query(F.data.startswith("w:nwiunit:"))
    async def cb_nwiunit(query: CallbackQuery) -> None:
        chat_id = query.message.chat.id  # type: ignore[union-attr]
        unit = query.data.split(":", 2)[2]  # type: ignore[union-attr]
        pending = _pending.get(chat_id)
        if not pending:
            await query.answer("Session expired.", show_alert=True)
            return
        label = _UNIT_LABELS[unit]
        await query.message.delete()  # type: ignore[union-attr]
        ask = await query.message.answer(  # type: ignore[union-attr]
            f"Send the interval in {label.lower()}:",
            reply_markup=_cancel_new_watcher_kb(),
        )
        _pending[chat_id] = {
            "action": "new_watcher_interval",
            "url": pending["url"],
            "name": pending["name"],
            "unit": unit,
            "ask_msg_id": ask.message_id,
        }
        await query.answer()

    @dp.callback_query(F.data == "w:skip_prompt")
    async def cb_skip_prompt(query: CallbackQuery, bot: Bot) -> None:
        chat_id = query.message.chat.id  # type: ignore[union-attr]
        pending = _pending.pop(chat_id, None)
        with contextlib.suppress(Exception):
            await query.message.delete()  # type: ignore[union-attr]
        await query.answer()
        if not pending:
            return
        new_w = wc.WatcherConfig(
            id=secrets.token_hex(4),
            name=pending["name"],
            url=pending["url"],
            interval=pending["interval"],
            enabled=True,
            created_at=datetime.now(timezone.utc).isoformat(),
            prompts=[],
        )
        wc.save(new_w)
        log.info("new watcher created id=%s name=%r url=%r", new_w.id, new_w.name, new_w.url)
        await bot.send_message(
            chat_id,
            _watcher_info_text(new_w),
            reply_markup=_actions_kb(new_w),
            parse_mode="HTML",
        )

    # ── Single message handler — dispatches by pending action ──────────────────

    @dp.message()
    async def handle_message(message: Message, bot: Bot) -> None:
        pending = _pending.get(message.chat.id)
        text = (message.text or "").strip()

        if pending is None:
            log.warning("unhandled message chat_id=%d user=%s text=%r",
                        message.chat.id,
                        message.from_user.username if message.from_user else None,
                        text)
            return

        action = pending["action"]
        ask_msg_id: int | None = pending.get("ask_msg_id")

        async def _cleanup_input() -> None:
            """Delete the bot's ask message and the user's reply."""
            if ask_msg_id:
                with contextlib.suppress(Exception):
                    await bot.delete_message(message.chat.id, ask_msg_id)
            with contextlib.suppress(Exception):
                await message.delete()

        # ── waiting for a new name ─────────────────────────────────────────────
        if action == "edit_name":
            if not text:
                await message.answer("Please send a non-empty name.")
                return
            w = wc.get(pending["watcher_id"])
            if w is None:
                _pending.pop(message.chat.id, None)
                await _cleanup_input()
                return
            w.name = text
            wc.save(w)
            _pending.pop(message.chat.id, None)
            await _cleanup_input()
            await bot.send_message(
                message.chat.id,
                f"{_watcher_info_text(w)}\n✅ Name updated.",
                reply_markup=_actions_kb(w),
                parse_mode="HTML",
            )

        # ── interval unit chosen via button — text input not expected ────────────
        elif action == "edit_interval_unit":
            await message.answer("Please choose a unit using the buttons above.")

        # ── waiting for a new interval value ───────────────────────────────────
        elif action == "edit_interval":
            if not text.isdigit() or int(text) <= 0:
                await message.answer("Please send a positive integer.")
                return
            w = wc.get(pending["watcher_id"])
            if w is None:
                _pending.pop(message.chat.id, None)
                await _cleanup_input()
                return
            unit = pending.get("unit", "s")
            w.interval = int(text) * _UNIT_MULT[unit]
            wc.save(w)
            _pending.pop(message.chat.id, None)
            await _cleanup_input()
            await bot.send_message(
                message.chat.id,
                f"{_watcher_info_text(w)}\n✅ Interval updated to <b>{w.interval}s</b>.",
                reply_markup=_actions_kb(w),
                parse_mode="HTML",
            )

        # ── waiting for replacement text for a specific prompt ─────────────────
        elif action == "edit_prompt_item":
            if not text:
                await message.answer("Please send non-empty text.")
                return
            wid = pending["watcher_id"]
            idx = pending["prompt_idx"]
            w = wc.get(wid)
            if w is None or idx >= len(w.prompts):
                _pending.pop(message.chat.id, None)
                await _cleanup_input()
                return
            w.prompts[idx] = text
            wc.save(w)
            _pending.pop(message.chat.id, None)
            await _cleanup_input()
            await _render_prompts(bot, message.chat.id, wid)

        # ── waiting for clipboard text ─────────────────────────────────────────
        elif action == "clipboard":
            if not text:
                await message.answer("Please send non-empty text.")
                return
            _pending.pop(message.chat.id, None)
            await _cleanup_input()
            ok = await _set_clipboard(text)
            if ok:
                await bot.send_message(message.chat.id, f"✅ Clipboard set via: {', '.join(ok)}")
            else:
                await bot.send_message(message.chat.id, "❌ No clipboard backend available (clip.exe / xclip / xsel).")

        # ── waiting for new prompt text to append ──────────────────────────────
        elif action == "create_prompt":
            if not text:
                await message.answer("Please send non-empty text.")
                return
            wid = pending["watcher_id"]
            w = wc.get(wid)
            if w is None:
                _pending.pop(message.chat.id, None)
                await _cleanup_input()
                return
            w.prompts.append(text)
            wc.save(w)
            _pending.pop(message.chat.id, None)
            await _cleanup_input()
            await _render_prompts(bot, message.chat.id, wid)

        # ── new watcher: waiting for URL ────────────────────────────────────────
        elif action == "new_watcher_url":
            if not text:
                await message.answer("Please send a non-empty URL.")
                return
            suggested = _name_from_url(text)
            await _cleanup_input()
            ask = await bot.send_message(
                message.chat.id,
                f"Send a name for this watcher (suggested: <code>{_html.escape(suggested)}</code>):",
                reply_markup=_cancel_new_watcher_kb(),
                parse_mode="HTML",
            )
            _pending[message.chat.id] = {
                "action": "new_watcher_name",
                "url": text,
                "ask_msg_id": ask.message_id,
            }

        # ── new watcher: waiting for name ───────────────────────────────────────
        elif action == "new_watcher_name":
            if not text:
                await message.answer("Please send a non-empty name.")
                return
            await _cleanup_input()
            ask = await bot.send_message(
                message.chat.id,
                "Choose the interval unit:",
                reply_markup=_unit_sel_new_kb(),
            )
            _pending[message.chat.id] = {
                "action": "new_watcher_interval_unit",
                "url": pending["url"],
                "name": text,
                "ask_msg_id": ask.message_id,
            }

        # ── new watcher: unit chosen via button — text input not expected ───────
        elif action == "new_watcher_interval_unit":
            await message.answer("Please choose a unit using the buttons above.")

        # ── new watcher: waiting for interval value ─────────────────────────────
        elif action == "new_watcher_interval":
            if not text.isdigit() or int(text) <= 0:
                await message.answer("Please send a positive integer.")
                return
            unit = pending.get("unit", "s")
            interval = int(text) * _UNIT_MULT[unit]
            await _cleanup_input()
            ask = await bot.send_message(
                message.chat.id,
                "Send the first prompt for this watcher, or tap Skip to add later:",
                reply_markup=_skip_prompt_kb(),
            )
            _pending[message.chat.id] = {
                "action": "new_watcher_prompt",
                "url": pending["url"],
                "name": pending["name"],
                "interval": interval,
                "ask_msg_id": ask.message_id,
            }

        # ── new watcher: waiting for first prompt ──────────────────────────────
        elif action == "new_watcher_prompt":
            if not text:
                await message.answer("Please send non-empty text, or tap Skip.")
                return
            await _cleanup_input()
            _pending.pop(message.chat.id, None)
            new_w = wc.WatcherConfig(
                id=secrets.token_hex(4),
                name=pending["name"],
                url=pending["url"],
                interval=pending["interval"],
                enabled=True,
                created_at=datetime.now(timezone.utc).isoformat(),
                prompts=[text],
            )
            wc.save(new_w)
            log.info("new watcher created id=%s name=%r url=%r", new_w.id, new_w.name, new_w.url)
            await bot.send_message(
                message.chat.id,
                _watcher_info_text(new_w),
                reply_markup=_actions_kb(new_w),
                parse_mode="HTML",
            )

    return dp


_BOT_COMMANDS = [
    BotCommand(command="watchers",  description="Watchers"),
    BotCommand(command="files",     description="Files"),
    BotCommand(command="clipboard", description="Clipboard"),
    BotCommand(command="pdf2docx",  description="PDF → DOCX"),
    BotCommand(command="status",    description="Status"),
    BotCommand(command="help",      description="Help"),
    BotCommand(command="start",     description="Start"),
]


async def run_bot(settings: Settings) -> None:
    """Start the bot and block until cancelled."""
    bot = Bot(token=settings.telegram.token)
    dp = _build_dispatcher(settings.telegram.chat_id, settings)

    await bot.set_my_commands(_BOT_COMMANDS)
    log.info("bot commands registered: %s", [c.command for c in _BOT_COMMANDS])

    log.info(
        "[bot] starting (chat_id=%d, poll_timeout=%ds)",
        settings.telegram.chat_id,
        settings.telegram.poll_timeout,
    )
    try:
        await dp.start_polling(bot, polling_timeout=settings.telegram.poll_timeout)
        log.info("[bot] start_polling returned normally")
    except asyncio.CancelledError:
        log.info("[bot] start_polling cancelled — closing session")
        raise
    except Exception:
        log.exception("[bot] start_polling raised unexpected exception")
        raise
    finally:
        log.info("[bot] closing bot session...")
        await bot.session.close()
        log.info("[bot] session closed")
